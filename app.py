import streamlit as st
from docx import Document
from datetime import datetime
import os

TEMPLATE_FILE = "offer_template.docx"

# Indian number system (Lakh, Crore)
ONES = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine"]
TENS = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
TEENS = ["Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen", "Seventeen", "Eighteen", "Nineteen"]

def _to_words_upto_99(n):
    if n == 0:
        return ""
    if n < 10:
        return ONES[n]
    if n < 20:
        return TEENS[n - 10]
    t, o = divmod(n, 10)
    return (TENS[t] + (" " + ONES[o] if o else "")).strip()

def _to_words_upto_999(n):
    if n == 0:
        return ""
    if n < 100:
        return _to_words_upto_99(n)
    h, r = divmod(n, 100)
    return (ONES[h] + " Hundred" + (" " + _to_words_upto_99(r) if r else "")).strip()

def number_to_words(num):
    """Convert number to words (Indian style: Lakh, Crore). E.g. 500000 -> 'Five Lakh Only'."""
    num = int(round(num))
    if num == 0:
        return "Zero Only"
    if num < 0:
        return "Minus " + number_to_words(-num)
    words = []
    # Crore
    crore, num = divmod(num, 10000000)
    if crore:
        words.append(_to_words_upto_999(crore) + " Crore")
    # Lakh
    lakh, num = divmod(num, 100000)
    if lakh:
        words.append(_to_words_upto_999(lakh) + " Lakh")
    # Thousand
    thousand, num = divmod(num, 1000)
    if thousand:
        words.append(_to_words_upto_999(thousand) + " Thousand")
    # Rest
    if num:
        words.append(_to_words_upto_999(num))
    return " ".join(words) + " Only"

def format_currency(num):
    return "{:,.0f}".format(num)

def replace_text_in_doc(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

st.title("Offer Letter Generator")

name = st.text_input("Candidate Name")
job_title = st.text_input("Job Title")
joining_deadline_date = st.date_input("Joining Deadline")
total_ctc = st.number_input("Total CTC", min_value=0)
st.caption(f"In words: **{number_to_words(total_ctc)}**")
performance_percent = st.number_input("Performance %", min_value=0.0)
probation = st.text_input("Probation Period", value="6")
notice = st.text_input("Notice Period", value="2")

if st.button("Generate Offer Letter"):

    performance_amount = total_ctc * performance_percent / 100
    fixed_ctc = total_ctc - performance_amount

    today = datetime.today().strftime("%d %B %Y")
    joining_deadline = joining_deadline_date.strftime("%d %B %Y")
    first_name = name.strip().split()[0] if name and name.strip() else name or ""

    replacements = {
        "{{date}}": today,
        "{{first_name}}": first_name,
        "{{full_name}}": name,
        "{{name}}": name,
        "{{job_title}}": job_title,
        "{{joining_deadline}}": joining_deadline,
        "{{total_ctc}}": format_currency(total_ctc),
        "{{total_ctc_words}}": number_to_words(total_ctc),
        "{{fixed_ctc}}": format_currency(fixed_ctc),
        "{{performance_percent}}": str(performance_percent),
        "{{performance_amount}}": format_currency(performance_amount),
        "{{probation}}": probation,
        "{{notice_period}}": notice,
    }

    doc = Document(TEMPLATE_FILE)
    replace_text_in_doc(doc, replacements)

    output_docx = f"./Offer_Letter_{name}-windmark.docx"
    doc.save(output_docx)

    st.success("Offer Letter Generated Successfully")

    with open(output_docx, "rb") as file:
        st.download_button("Download Word", file, file_name=os.path.basename(output_docx))